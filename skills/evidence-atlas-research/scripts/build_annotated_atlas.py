#!/usr/bin/env python3
"""Build a two-up Chinese patent/mechanism atlas DOCX from a JSON manifest.

Manifest schema:
{
  "title": "仿生鳄鱼机器人",
  "subtitle": "专利结构图解 · 清晰注释版",
  "footer": "仿生鳄鱼图册",
  "cover_note": "optional",
  "closing_note": "optional",
  "sources": [{
    "id": "CN...", "title": "...", "status": "...", "url": "https://...",
    "figures": [{
      "image": "path/to/image.png", "label": "图纸页 1",
      "what": "...", "motion": "...", "prototype": "..."
    }]
  }]
}

Relative image paths resolve against the manifest directory.
"""

from __future__ import annotations

import argparse
import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BLUE, LIGHT, PALE, GRAY, INK, GOLD, RED = (
    "1F4E79", "EAF1F8", "F5F8FB", "5A6570", "17212B", "9A6A00", "9B1C1C"
)


def font(run, size=10, bold=False, color=INK):
    name = "Microsoft YaHei"
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    for key in ("ascii", "hAnsi", "eastAsia"):
        rpr.rFonts.set(qn(f"w:{key}"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    node = tcpr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tcpr.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell, top=110, start=120, bottom=110, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def geometry(table, widths=(4680, 4680), indent=120):
    table.autofit = False
    pr = table._tbl.tblPr
    tw = pr.find(qn("w:tblW"))
    if tw is None:
        tw = OxmlElement("w:tblW")
        pr.append(tw)
    tw.set(qn("w:w"), str(sum(widths)))
    tw.set(qn("w:type"), "dxa")
    ti = pr.find(qn("w:tblInd"))
    if ti is None:
        ti = OxmlElement("w:tblInd")
        pr.append(ti)
    ti.set(qn("w:w"), str(indent))
    ti.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")


def no_borders(table):
    pr = table._tbl.tblPr
    borders = pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)


def hyperlink(paragraph, text, url):
    rid = paragraph.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    for tag, attr, value in (("color", "val", BLUE), ("u", "val", "single"), ("sz", "val", "18")):
        node = OxmlElement(f"w:{tag}")
        node.set(qn(f"w:{attr}"), value)
        rpr.append(node)
    run.append(rpr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText"); text.set(qn("xml:space"), "preserve"); text.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, end])


def note(doc, text, fill=LIGHT, color=INK):
    table = doc.add_table(rows=1, cols=1)
    geometry(table, (9360,))
    cell = table.cell(0, 0)
    shade(cell, fill)
    margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(text), 9.1, color=color)


def heading(doc, text, size=15):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(text), size, True, BLUE)


def panel(cell, image, figure):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    margins(cell, 120, 130, 120, 130)
    shade(cell, PALE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(image), width=Inches(2.93))
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(f"{figure['label']}｜{figure['what']}"), 10, True, BLUE)
    for label, key in (("运动/作用：", "motion"), ("样机借鉴：", "prototype")):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.12
        font(p.add_run(label), 8.6, True)
        font(p.add_run(figure[key]), 8.6)


def validate(data, base):
    for key in ("title", "sources"):
        if key not in data:
            raise ValueError(f"manifest missing {key}")
    if not data["sources"]:
        raise ValueError("manifest has no sources")
    for source in data["sources"]:
        for key in ("id", "title", "url", "figures"):
            if not source.get(key):
                raise ValueError(f"source missing {key}: {source.get('id', '?')}")
        for index, figure in enumerate(source["figures"], 1):
            for key in ("image", "label", "what", "motion", "prototype"):
                if not figure.get(key):
                    raise ValueError(f"{source['id']} figure {index} missing {key}")
            path = Path(figure["image"])
            if not path.is_absolute():
                path = base / path
            if not path.is_file():
                raise ValueError(f"image not found: {path}")


def build(data, manifest_path, output):
    base = manifest_path.parent
    validate(data, base)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = data.get("footer", data["title"])
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.runs[0], 8, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run(f"{data.get('footer', data['title'])}  |  "), 8, color=GRAY)
    page_field(footer)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(110)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run(data["title"]), 30, True, BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    font(p.add_run(data.get("subtitle", "结构图解 · 清晰注释版")), 17, True)
    count = sum(len(source["figures"]) for source in data["sources"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(44)
    font(p.add_run(f"{len(data['sources'])} 组来源 · {count} 张图纸页 · 每图三句说明"), 11, color=GRAY)
    note(doc, data.get("cover_note", "读图规则：原图保持不变；图示和运动来自来源文本与图面；样机借鉴是工程建议。"))
    doc.add_page_break()

    heading(doc, "来源目录", 22)
    for source in data["sources"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        font(p.add_run(f"{source['id']}  "), 11, True, BLUE)
        font(p.add_run(source["title"]), 10.5, True)
        font(p.add_run(f"  ·  {len(source['figures'])} 张图纸页"), 9, color=GRAY)
    doc.add_page_break()

    for source_index, source in enumerate(data["sources"]):
        figures = source["figures"]
        for start in range(0, len(figures), 2):
            heading(doc, f"{source['id']}｜{source['title']}", 14.3)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            if source.get("status"):
                font(p.add_run(f"状态线索：{source['status']}  ·  "), 8.7, color=GRAY)
            hyperlink(p, "打开原始来源", source["url"])
            table = doc.add_table(rows=1, cols=2)
            geometry(table)
            no_borders(table)
            for offset in range(2):
                index = start + offset
                cell = table.cell(0, offset)
                if index >= len(figures):
                    shade(cell, "FFFFFF")
                    cell.text = ""
                    continue
                figure = figures[index]
                image = Path(figure["image"])
                if not image.is_absolute():
                    image = base / image
                panel(cell, image, figure)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            font(p.add_run(f"原始图未重绘 · 图组 {start // 2 + 1}/{(len(figures) + 1) // 2}"), 7.8, color=GRAY)
            last_source = source_index == len(data["sources"]) - 1
            last_pair = start + 2 >= len(figures)
            if not (last_source and last_pair):
                doc.add_page_break()

    if data.get("closing_note"):
        note(doc, data["closing_note"], "FCE8E6", RED)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    with zipfile.ZipFile(output) as archive:
        bad = archive.testzip()
    if bad:
        raise RuntimeError(f"corrupt DOCX member: {bad}")


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("manifest", type=Path)
    parser.add_argument("--out", required=True, type=Path)
    args = parser.parse_args()
    try:
        data = json.loads(args.manifest.read_text(encoding="utf-8"))
        build(data, args.manifest.resolve(), args.out.resolve())
    except Exception as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 1
    print(args.out.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
